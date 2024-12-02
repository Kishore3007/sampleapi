import os
import PyPDF2
import docx
import pythoncom
from win32com.client import Dispatch
from fastapi import HTTPException

# Function to extract text from a PDF
def extract_pdf_text(file_path):
    text = ""
    try:
        with open(file_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                text += f"{page_text}\n"
                print(text)
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {e}")

# Function to extract text from a DOCX file
def extract_docx_text(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Avoid empty paragraphs
                full_text.append(para.text)

        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(' | '.join(row_text))  # Format table rows with a separator

        return "\n".join(full_text) if full_text else None
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from DOCX: {e}")

# Function to extract text from a DOC file using COM (Windows only)
def extract_doc_text(file_path):
    try:
        pythoncom.CoInitialize()  # Initialize the COM library for thread support
        
        # Convert the file path to an absolute path with double backslashes
        absolute_file_path = os.path.abspath(file_path)
        absolute_file_path = absolute_file_path.replace("/", "\\")  # Convert to backslashes for Windows

        # Dispatch the Word application and open the document
        word_app = Dispatch('Word.Application')
        word_app.Visible = False  # Optional: Run Word in the background
        doc = word_app.Documents.Open(absolute_file_path)
        doc_text = doc.Content.Text  # Extract the text from the document

        # Close the document and Word application
        doc.Close()
        word_app.Quit()

        return doc_text.strip()  # Return the cleaned text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from DOC: {e}")
    finally:
        pythoncom.CoUninitialize()  # Uninitialize the COM library
